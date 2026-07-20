import os
import io
from datetime import datetime
from celery import Celery
from unstructured.partition.auto import partition
import docx
import csv
import json

# Initialize Celery app
celery_app = Celery(
    "tasks",
    broker=os.getenv("REDIS_URL", "redis://127.0.0.1:6379/0"),
    backend=os.getenv("REDIS_URL", "redis://127.0.0.1:6379/0")
)

# Tasks queue settings
celery_app.conf.update(
    task_serializer="json",
    accept_content=["json"],
    result_serializer="json",
    timezone="UTC",
    enable_utc=True,
)


def get_file_metadata(file_path: str, content: str = "") -> dict:
    """
    Extracts document metadata: file size, estimated page count, creation timestamp, and author.
    """
    try:
        stat_info = os.stat(file_path)
        file_size_kb = round(stat_info.st_size / 1024, 2)
        creation_date = datetime.fromtimestamp(stat_info.st_ctime).isoformat()
    except Exception:
        file_size_kb = 0.0
        creation_date = datetime.now().isoformat()

    # Estimate page count from content length (~3000 chars per page)
    page_count = max(1, len(content) // 3000) if content else 1

    # Try to extract author from .docx core properties
    author = "Unknown"
    if file_path.lower().endswith(".docx"):
        try:
            doc = docx.Document(file_path)
            author = doc.core_properties.author or "Unknown"
        except Exception:
            pass

    return {
        "author": author,
        "creation_date": creation_date,
        "file_size_kb": file_size_kb,
        "page_count": page_count
    }


@celery_app.task(name="tasks.ingest_file")
def ingest_file_task(file_path: str, filename: str, username: str):
    """
    Celery background task to parse multi-format files in a layout-aware manner.
    Returns the raw parsed content for the FastAPI server to embed and index.
    """
    try:
        content = parse_file_content(file_path)
    except Exception as e:
        print(f"[ERROR] Failed to parse {filename}: {e}")
        return {"status": "failed", "error": str(e)}

    if not content.strip():
        return {"status": "failed", "error": "Empty document content"}

    # Extract file metadata for storage in vector DB payload
    doc_metadata = get_file_metadata(file_path, content)

    return {
        "status": "completed",
        "content": content,
        "filename": filename,
        "username": username,
        "file_path": file_path,
        "doc_metadata": doc_metadata
    }


def parse_file_content(file_path: str) -> str:
    """
    Robust multi-format local document parser with fallback capabilities.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    # 1. Word document parsing
    if ext == ".docx":
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
        
    # 2. CSV document parsing
    elif ext == ".csv":
        rows = []
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))
        return "\n".join(rows)
        
    # 3. JSON parsing
    elif ext == ".json":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            data = json.load(f)
            return json.dumps(data, indent=2)
            
    # 4. PDF Layout-aware extraction using Unstructured, fallback to PyMuPDF OCR
    elif ext == ".pdf":
        try:
            elements = partition(filename=file_path)
            text = "\n".join([str(el) for el in elements])
            if len(text.strip()) > 100:
                return text
        except Exception:
            pass
            
        # Fallback to PyMuPDF + Tesseract OCR
        # Import lazily to avoid loading heavy modules unless PDF fallback is triggered
        import fitz
        from PIL import Image
        import pytesseract
        
        content = ""
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text = page.get_text()
                # Run OCR if page is image-only
                if not text.strip() or len(text.strip()) < 15:
                    pix = page.get_pixmap(dpi=150)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    text = pytesseract.image_to_string(img)
                if text:
                    content += text + "\n"
        except Exception as e:
            print(f"[OCR FALLBACK ERROR] {e}")
            
        return content
        
    # 5. Plain Text / Markdown
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

